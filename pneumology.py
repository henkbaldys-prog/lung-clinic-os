# pneumology.py - Lung Clinic OS
# Complete standalone Pneumology Application
# 10 tabs: Intake, Spirometry, COPD/Asthma, Imaging, Nodule/Cancer,
#          Infection, Respiratory Failure/ICU, Sleep Medicine, Follow-up, Discharge
#
# Workflow: dictation / notes -> AI extracts -> ready clinical document
# No patient-identifying data. GDPR / HIPAA safe. Anonymous only.
# Languages: de / en / ro / es / ar
#
# STANDALONE USE - run as main app.py or call render() from your main app:
#
#   import pneumology
#   pneumology.render(lang_code, client, supabase, user_code)
#
# Or as standalone:
#   streamlit run pneumology.py

import streamlit as st
import streamlit.components.v1 as components
from openai import OpenAI
from datetime import date, datetime
import os, io, tempfile, subprocess, math

try:
    from docx import Document as DocxDoc
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import imageio_ffmpeg as _ffmpeg
    FFMPEG_PATH = _ffmpeg.get_ffmpeg_exe()
except Exception:
    FFMPEG_PATH = None

# ---------------------------------------------------------------------------
# TRANSLATIONS
# ---------------------------------------------------------------------------
TX = {
    "app_title":   {"de":"Lung Clinic OS","en":"Lung Clinic OS","ro":"Lung Clinic OS","es":"Lung Clinic OS","ar":"Lung Clinic OS"},
    "app_sub":     {"de":"Vollstaendiges Pneumologie-Dokumentationssystem","en":"Complete Pneumology Documentation System","ro":"Sistem Complet de Documentatie Pneumologie","es":"Sistema Completo de Documentacion Neumologia","ar":"    "},
    "disclaimer":  {"de":"Nur Dokumentation. Keine automatische Diagnose. Arzt entscheidet. Keine patientenidentifizierenden Daten eingeben.","en":"Documentation only. No automatic diagnosis. Physician decides. Do not enter patient-identifying data.","ro":"Doar documentare. Fara diagnostic automat. Medicul decide. Nu introduceti date identificatoare.","es":"Solo documentacion. Sin diagnostico automatico. El medico decide. No introduzca datos identificativos.","ar":" .   .  .     ."},
    # Tabs
    "t_intake":    {"de":"Aufnahme","en":"Intake","ro":"Internare","es":"Ingreso","ar":""},
    "t_spiro":     {"de":"Spirometrie / PFT","en":"Spirometry / PFT","ro":"Spirometrie / PFT","es":"Espirometria / PFT","ar":" "},
    "t_copd":      {"de":"COPD / Asthma","en":"COPD / Asthma","ro":"BPOC / Astm","es":"EPOC / Asma","ar":"COPD / "},
    "t_imaging":   {"de":"Bildgebung","en":"Imaging","ro":"Imagistica","es":"Imagen","ar":""},
    "t_nodule":    {"de":"Nodulus / Tumor","en":"Nodule / Cancer","ro":"Nodul / Cancer","es":"Nodulo / Cancer","ar":" / "},
    "t_infection": {"de":"Lungeninfektion","en":"Lung Infection","ro":"Infectie Pulmonara","es":"Infeccion Pulmonar","ar":" "},
    "t_icu":       {"de":"Atemversagen / ICU","en":"Resp. Failure / ICU","ro":"Insuf. Respiratorie / ATI","es":"Insuf. Respiratoria / UCI","ar":"  / ICU"},
    "t_sleep":     {"de":"Schlafmedizin","en":"Sleep Medicine","ro":"Medicina Somnului","es":"Medicina del Sueno","ar":" "},
    "t_followup":  {"de":"Verlauf / Monitoring","en":"Follow-up / Monitoring","ro":"Evolutie / Monitorizare","es":"Seguimiento / Monitoreo","ar":" / "},
    "t_discharge": {"de":"Entlassung","en":"Discharge","ro":"Externare","es":"Alta","ar":""},
    # Common UI
    "notes_label": {"de":"Klinische Notizen / Diktat (anonym)","en":"Clinical notes / dictation (anonymous)","ro":"Note clinice / dictare (anonim)","es":"Notas clinicas / dictado (anonimo)","ar":" /  ()"},
    "notes_hint":  {"de":"Eintippen oder diktieren. KEIN Name, KEIN Geburtsdatum, KEIN Datum.","en":"Type or dictate. NO name, NO date of birth, NO date.","ro":"Tastati sau dictati. FARA nume, FARA data nasterii.","es":"Escriba o dicte. SIN nombre, SIN fecha nacimiento.","ar":"  .     ."},
    "audio_label": {"de":"Sprachdatei hochladen (Whisper)","en":"Upload audio file (Whisper)","ro":"Incarcati fisier audio (Whisper)","es":"Subir audio (Whisper)","ar":"   (Whisper)"},
    "transcribed": {"de":"Transkription OK","en":"Transcription OK","ro":"Transcriere OK","es":"Transcripcion OK","ar":" "},
    "trans_err":   {"de":"Transkriptionsfehler","en":"Transcription error","ro":"Eroare transcriere","es":"Error transcripcion","ar":"  "},
    "generate":    {"de":"Bericht generieren","en":"Generate Report","ro":"Genereaza Raport","es":"Generar Informe","ar":" "},
    "generating":  {"de":"Generiere...","en":"Generating...","ro":"Se genereaza...","es":"Generando...","ar":" ..."},
    "copy":        {"de":"Kopieren","en":"Copy","ro":"Copiaza","es":"Copiar","ar":""},
    "dl_txt":      {"de":"Download .txt","en":"Download .txt","ro":"Descarca .txt","es":"Descargar .txt","ar":" .txt"},
    "dl_docx":     {"de":"Download .docx","en":"Download .docx","ro":"Descarca .docx","es":"Descargar .docx","ar":" .docx"},
    "new":         {"de":"Neu","en":"New","ro":"Nou","es":"Nuevo","ar":""},
    "fill_warn":   {"de":"Bitte Notizen eingeben oder Audio hochladen.","en":"Please enter notes or upload audio.","ro":"Introduceti note sau incarcati audio.","es":"Introduzca notas o suba audio.","ar":"      ."},
    "no_id":       {"de":"Keine Patientendaten (DSGVO). Anonym.","en":"No patient data (GDPR). Anonymous.","ro":"Fara date pacient (GDPR). Anonim.","es":"Sin datos paciente (GDPR). Anonimo.","ar":"   (GDPR). ."},
    "insuf_data":  {"de":"Unzureichende Daten - bitte mehr Informationen eingeben.","en":"Insufficient data - please enter more information.","ro":"Date insuficiente - introduceti mai multe informatii.","es":"Datos insuficientes - introduzca mas informacion.","ar":"   -     ."},
    # Intake
    "main_symptom":{"de":"Hauptsymptom","en":"Main Symptom","ro":"Simptom Principal","es":"Sintoma Principal","ar":" "},
    "symp_opts":   {
        "de":["Dyspnoe / Atemnot","Husten (akut)","Husten (chronisch)","Haemoptyse / Bluthusten","Brustschmerz / -enge","Giemen / Pfeifen","Schlafapnoe-Verdacht","Thoraxschmerz pleuritisch","Gewichtsverlust / B-Symptomatik","Stridor","Sonstige"],
        "en":["Dyspnea / Shortness of breath","Cough (acute)","Cough (chronic)","Hemoptysis","Chest tightness / pain","Wheeze","Suspected sleep apnea","Pleuritic chest pain","Weight loss / B-symptoms","Stridor","Other"],
        "ro":["Dispnee","Tuse (acuta)","Tuse (cronica)","Hemoptizie","Durere / presiune toracica","Wheezing","Apnee in somn suspect","Durere pleuritica","Scadere ponderala / Simptome B","Stridor","Altele"],
        "es":["Disnea","Tos aguda","Tos cronica","Hemoptisis","Opresion / dolor toracico","Sibilancias","Apnea del sueno sospechada","Dolor pleuritico","Perdida peso / sintomas B","Estridor","Otro"],
        "ar":[" "," "," "," "," /  ","","   "," ","  /  B","",""],
    },
    "smoking":     {"de":"Raucherstatus","en":"Smoking status","ro":"Statut fumat","es":"Estado tabaquismo","ar":" "},
    "smoking_opts":{"de":["Nie geraucht","Aktuell Raucher","Ex-Raucher","Passivrauch-Exposition"],"en":["Never","Current smoker","Ex-smoker","Passive exposure"],"ro":["Nefumator","Fumator activ","Fost fumator","Expunere pasiva"],"es":["Nunca","Fumador activo","Ex-fumador","Exposicion pasiva"],"ar":[""," ",""," "]},
    "pack_years":  {"de":"Packungsjahre (pack-years)","en":"Pack-years","ro":"Pachete-ani","es":"Paquetes-ano","ar":" "},
    "occupation":  {"de":"Berufsanamnese (Asbest, Staub, Chemikalien)","en":"Occupational history (asbestos, dust, chemicals)","ro":"Expunere profesionala (azbest, praf, chimicale)","es":"Historia laboral (asbesto, polvo, quimicos)","ar":"  (   )"},
    "country":     {"de":"Land / System","en":"Country / System","ro":"Tara / Sistem","es":"Pais / Sistema","ar":" / "},
    "country_opts":{"de":["Deutschland (ICD-10)","Oesterreich","Schweiz","USA (CPT/ICD-10-CM)","Saudi-Arabien","VAE","Rumaenien","Spanien","Sonstiges"],"en":["Germany (ICD-10)","Austria","Switzerland","USA (CPT/ICD-10-CM)","Saudi Arabia","UAE","Romania","Spain","Other"],"ro":["Germania","Austria","Elvetia","SUA (CPT)","Arabia Saudita","EAU","Romania","Spania","Altele"],"es":["Alemania (ICD-10)","Austria","Suiza","EE.UU. (CPT)","Arabia Saudi","EAU","Rumania","Espana","Otro"],"ar":["","",""," ","","","","",""]},
    # COPD/Asthma
    "copd_stage":  {"de":"GOLD-Stadium (klinisch eingeschaetzt)","en":"GOLD Stage (clinically estimated)","ro":"Stadiu GOLD","es":"Estadio GOLD","ar":" GOLD"},
    "gold_opts":   {"de":["GOLD I (leicht)","GOLD II (mittel)","GOLD III (schwer)","GOLD IV (sehr schwer)","Unklar / noch nicht klassifiziert"],"en":["GOLD I (mild)","GOLD II (moderate)","GOLD III (severe)","GOLD IV (very severe)","Unclear / not yet classified"],"ro":["GOLD I (usor)","GOLD II (moderat)","GOLD III (sever)","GOLD IV (foarte sever)","Neclar"],"es":["GOLD I (leve)","GOLD II (moderado)","GOLD III (grave)","GOLD IV (muy grave)","No clasificado"],"ar":["GOLD I ()","GOLD II ()","GOLD III ()","GOLD IV ( )"," "]},
    "asthma_ctrl": {"de":"Asthmakontrolle","en":"Asthma control","ro":"Control astm","es":"Control asma","ar":" "},
    "asthma_opts": {"de":["Kontrolliert","Teilweise kontrolliert","Unkontrolliert"],"en":["Controlled","Partly controlled","Uncontrolled"],"ro":["Controlat","Partial controlat","Necontrolat"],"es":["Controlado","Parcialmente controlado","No controlado"],"ar":[""," "," "]},
    # Nodule
    "nodule_size": {"de":"Nodulus-Groesse (mm)","en":"Nodule size (mm)","ro":"Marime nodul (mm)","es":"Tamano nodulo (mm)","ar":"  ()"},
    "nodule_type": {"de":"Nodulus-Typ","en":"Nodule type","ro":"Tip nodul","es":"Tipo nodulo","ar":" "},
    "nodule_opts": {"de":["Solid","Teils-solid (part-solid)","Milchglas (ground-glass)","Multipel"],"en":["Solid","Part-solid","Ground-glass opacity","Multiple"],"ro":["Solid","Partial solid","Ground-glass","Multiple"],"es":["Solido","Parcialmente solido","Vidrio esmerilado","Multiple"],"ar":[""," "," ",""]},
    "nodule_growth":{"de":"Bekanntes Wachstum?","en":"Known growth?","ro":"Crestere cunoscuta?","es":"Crecimiento conocido?","ar":" "},
    # Infection
    "infect_type": {"de":"Infektionstyp","en":"Infection type","ro":"Tip infectie","es":"Tipo infeccion","ar":" "},
    "infect_opts": {"de":["Ambulant erworbene Pneumonie (CAP)","Nosokomial (HAP/VAP)","Aspiration","Tuberkulose (V.a.)","Virale Pneumonie","Pilzpneumonie (V.a.)","Pleuraempyem"],"en":["Community-acquired pneumonia (CAP)","Hospital-acquired (HAP/VAP)","Aspiration pneumonia","Tuberculosis (suspected)","Viral pneumonia","Fungal pneumonia (suspected)","Pleural empyema"],"ro":["Pneumonie comunitara (CAP)","Nozocomiala (HAP/VAP)","Aspiratie","Tuberculoza (suspect)","Pneumonie virala","Pneumonie fungica","Empiem pleural"],"es":["Neumonia comunitaria (NAC)","Nosocomial (NAH/VAP)","Aspiracion","Tuberculosis (sospecha)","Neumonia viral","Neumonia fungica","Empiema pleural"],"ar":["   (CAP)","  ",""," ()","  ","  "," "]},
    # ICU / Resp failure
    "resp_fail_type":{"de":"Typ des Atemversagens","en":"Type of respiratory failure","ro":"Tip insuficienta respiratorie","es":"Tipo insuficiencia respiratoria","ar":"  "},
    "rf_opts":     {"de":["Typ I (hypoxaemisch)","Typ II (hyperkapnisch)","Gemischt","Noch unklar"],"en":["Type I (hypoxemic)","Type II (hypercapnic)","Mixed","Unclear"],"ro":["Tip I (hipoxemic)","Tip II (hipercapnic)","Mixt","Neclar"],"es":["Tipo I (hipoxemico)","Tipo II (hipercapnico)","Mixto","No claro"],"ar":[" I ( )"," II ( CO2)",""," "]},
    "o2_support":  {"de":"O2-Therapie / Beatmung","en":"O2 support / Ventilation","ro":"Suport O2 / Ventilatie","es":"Soporte O2 / Ventilacion","ar":"  / "},
    "o2_opts":     {"de":["Raumluft (21%)","Nasale O2 (1-6 L/min)","High-flow Nasalkanule (HFNC)","Maske (einfach)","Reservoirmaske","NIV / CPAP / BiPAP","Endotracheale Intubation (ETT)"],"en":["Room air (21%)","Nasal cannula O2","HFNC","Simple mask","Non-rebreather mask","NIV / CPAP / BiPAP","Endotracheal intubation (ETT)"],"ro":["Aer ambient","O2 nazal","HFNC","Masca simpla","Masca cu rezervor","VNI / CPAP / BiPAP","Intubatie endotraheala (IOT)"],"es":["Aire ambiente","O2 nasal","HFNC","Mascarilla simple","Mascarilla con reservorio","VNI / CPAP / BiPAP","Intubacion endotraqueal"],"ar":[" ","  ","HFNC"," ","  ","NIV / CPAP / BiPAP","  "]},
    # Sleep
    "stopbang_label":{"de":"STOP-BANG Score","en":"STOP-BANG Score","ro":"STOP-BANG Score","es":"STOP-BANG Score","ar":" STOP-BANG"},
    # Discharge
    "dc_scenario": {"de":"Entlassungs-Szenario","en":"Discharge scenario","ro":"Scenariu externare","es":"Escenario alta","ar":" "},
    "dc_opts":     {"de":["COPD-Entlassung","Asthma-Entlassung","Pneumonie-Entlassung","Lungentumor / Staging","ICU-Entlassung / Beatmungsentwaehnung","Pleuraerkrankung","ILD / Fibrose","Sonstiges"],"en":["COPD discharge","Asthma discharge","Pneumonia discharge","Lung tumor / Staging","ICU / Ventilator weaning discharge","Pleural disease","ILD / Fibrosis","Other"],"ro":["Externare BPOC","Externare astm","Externare pneumonie","Tumor pulmonar / Stadializare","Externare ATI","Boala pleurala","FPI / ILD","Altele"],"es":["Alta EPOC","Alta asma","Alta neumonia","Tumor pulmonar / Estadificacion","Alta UCI / Destete","Enfermedad pleural","FPI / EPI","Otro"],"ar":[" COPD"," ","  ","  / "," ICU /   "," ","ILD / ",""]},
}


def tx(key, lang):
    return TX[key].get(lang, TX[key].get("en", key))


LANG_NAME = {"de":"German","en":"English","ro":"Romanian","es":"Spanish","ar":"Arabic"}

COUNTRY_STYLE = {
    "Deutschland (ICD-10)": "German pneumology style: formal, detailed, ICD-10 codes, ATS/ERS and GOLD/GINA guideline citations.",
    "Oesterreich": "Austrian German pneumology style: formal, ICD-10, detailed.",
    "Schweiz": "Swiss German pneumology style: formal, detailed.",
    "USA (CPT/ICD-10-CM)": "US pulmonology style: concise, ATS/ACCP guideline-based, CPT codes, liability-safe language.",
    "Saudi Arabia": "Formal Arabic-influenced English. Concise, respectful, professional.",
    "UAE": "Formal English. Professional and concise.",
    "Romania": "Romanian pneumology style: functional, structured, ICD-10.",
    "Spain": "Spanish pneumology style: clear, SEPAR guideline-based.",
    "Germany (ICD-10)": "German pneumology style: formal, detailed, ICD-10, ATS/ERS citations.",
}

# ---------------------------------------------------------------------------
# SESSION STATE
# ---------------------------------------------------------------------------
def _init():
    defaults = {
        "pn_country": 0, "pn_lang": "en",
        "pn_symptom": 0, "pn_smoking": 0, "pn_pack_years": 0.0,
        "pn_occupation": "",
        # Spirometry
        "pn_fev1": 0.0, "pn_fvc": 0.0, "pn_fev1_fvc": 0.0,
        "pn_fev1_pct": 0.0, "pn_fvc_pct": 0.0, "pn_dlco_pct": 0.0,
        # COPD/Asthma
        "pn_gold": 0, "pn_asthma_ctrl": 0,
        "pn_exac_history": 0,
        # Nodule
        "pn_nodule_size": 0.0, "pn_nodule_type": 0,
        "pn_nodule_growth": False, "pn_nodule_prev_size": 0.0,
        # Infection
        "pn_infect_type": 0,
        # CURB-65
        "pn_curb_conf": False, "pn_curb_urea": False,
        "pn_curb_rr": False, "pn_curb_bp": False, "pn_curb_age": False,
        # ICU
        "pn_rf_type": 0, "pn_spo2": 95, "pn_po2": 0.0,
        "pn_pco2": 0.0, "pn_rr_resp": 16, "pn_o2_support": 0,
        "pn_fio2": 21,
        # Sleep STOP-BANG
        "pn_sb_snore": False, "pn_sb_tired": False, "pn_sb_obs": False,
        "pn_sb_bp": False, "pn_sb_bmi": False, "pn_sb_age": False,
        "pn_sb_neck": False, "pn_sb_male": False,
        # Discharge
        "pn_dc_scenario": 0,
        # Results
        "pn_res_intake": "", "pn_res_spiro": "", "pn_res_copd": "",
        "pn_res_imaging": "", "pn_res_nodule": "", "pn_res_infection": "",
        "pn_res_icu": "", "pn_res_sleep": "", "pn_res_followup": "",
        "pn_res_discharge": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------
def _convert_audio(audio_bytes, ffmpeg_path):
    with tempfile.NamedTemporaryFile(suffix=".input", delete=False) as tmp_in:
        tmp_in.write(audio_bytes)
        tmp_in_path = tmp_in.name
    tmp_out = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
    tmp_out_path = tmp_out.name
    tmp_out.close()
    cmd = [ffmpeg_path, "-y", "-i", tmp_in_path,
           "-ar", "16000", "-ac", "1", "-f", "wav", tmp_out_path]
    result = subprocess.run(cmd, capture_output=True, text=True)
    os.remove(tmp_in_path)
    if result.returncode != 0:
        return None, result.stderr
    return tmp_out_path, None


def _audio_input(lang, key, openai_client):
    if openai_client is None:
        return ""
    uploaded = st.file_uploader(
        tx("audio_label", lang),
        type=["wav","mp3","m4a","mp4","mpeg","mpga","ogg","webm"],
        key="audio_" + key
    )
    if uploaded is not None:
        if FFMPEG_PATH is None:
            st.warning("FFmpeg not available - audio upload disabled.")
            return ""
        audio_bytes = uploaded.read()
        wav_path, err = _convert_audio(audio_bytes, FFMPEG_PATH)
        if err:
            st.error(tx("trans_err", lang) + ": " + err)
            return ""
        try:
            with open(wav_path, "rb") as f:
                trans = openai_client.audio.transcriptions.create(
                    model="whisper-1", file=f
                )
            os.remove(wav_path)
            st.success(tx("transcribed", lang))
            return trans.text
        except Exception as e:
            st.error(tx("trans_err", lang) + ": " + str(e))
            if os.path.exists(wav_path):
                os.remove(wav_path)
    return ""


def _notes_and_audio(lang, tab_key, openai_client, height=180, placeholder_extra=""):
    st.caption(tx("notes_hint", lang))
    st.info(tx("no_id", lang))
    ph = {
        "de":"Pneumologische Notizen eingeben oder diktieren..." + placeholder_extra,
        "en":"Enter pulmonary notes or dictate..." + placeholder_extra,
        "ro":"Introduceti note pneumologice sau dictati...",
        "es":"Introduzca notas neumologicas o dicte...",
        "ar":"    ..."
    }.get(lang, "Enter notes or dictate...")
    notes = st.text_area(
        tx("notes_label", lang),
        height=height,
        key="notes_" + tab_key,
        placeholder=ph
    )
    transcribed = _audio_input(lang, tab_key, openai_client)
    if transcribed:
        notes = (notes + "\n" + transcribed).strip()
    return notes


def _copy_button(text, key, lang):
    label = tx("copy", lang)
    escaped = (text
               .replace("\\", "\\\\")
               .replace("`", "\\`")
               .replace("$", "\\$")
               .replace("\n", "\\n"))
    components.html(
        "<button onclick=\"navigator.clipboard.writeText(`" + escaped + "`)"
        ".then(()=>{let b=document.getElementById('cpb_" + key + "');"
        "b.innerText='OK!';setTimeout(()=>b.innerText='" + label + "',2000);})\""
        " id=\"cpb_" + key + "\""
        " style=\"background:#0369a1;color:#fff;border:none;padding:8px 20px;"
        "border-radius:6px;cursor:pointer;font-size:14px;font-weight:600;"
        "font-family:system-ui,sans-serif;\">" + label + "</button>",
        height=46
    )


def _make_docx(text, title):
    doc = DocxDoc()
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("Generated: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
    doc.add_paragraph("-" * 50)
    for line in text.split("\n"):
        doc.add_paragraph(line.strip())
    p = doc.add_paragraph()
    p.add_run("Lung Clinic OS - Documentation only. Physician decides all clinical actions.").italic = True
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _output_block(result_key, title, lang):
    text = st.session_state.get(result_key, "")
    if not text:
        return
    st.markdown("---")
    st.text_area(title, value=text, height=340, key="out_" + result_key)
    c1, c2, c3 = st.columns([2, 1, 1])
    with c1:
        _copy_button(text, result_key, lang)
    with c2:
        st.download_button(
            tx("dl_txt", lang),
            data=text.encode("utf-8"),
            file_name=result_key + "_" + str(date.today()) + ".txt",
            mime="text/plain",
            use_container_width=True
        )
    with c3:
        if DOCX_OK:
            st.download_button(
                tx("dl_docx", lang),
                data=_make_docx(text, title),
                file_name=result_key + "_" + str(date.today()) + ".docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.caption("pip install python-docx")


def _call_ai(prompt, openai_client, max_tokens=2000):
    if openai_client is None:
        return "OpenAI client not configured."
    resp = openai_client.chat.completions.create(
        model="gpt-4.1",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=max_tokens
    )
    return resp.choices[0].message.content


def _country_info(lang, idx):
    en_opts = ["Germany (ICD-10)","Austria","Switzerland","USA (CPT/ICD-10-CM)",
               "Saudi Arabia","UAE","Romania","Spain","Other"]
    country_en = en_opts[min(idx, len(en_opts)-1)]
    style = COUNTRY_STYLE.get(country_en, "Write professionally in " + LANG_NAME.get(lang,"the selected language") + ".")
    c_opts = tx("country_opts", lang)
    return c_opts[min(idx, len(c_opts)-1)], style


RULES = (
    "RULES:\n"
    "- Write in {lang}. {style}\n"
    "- Extract ALL data from notes. Do NOT invent values.\n"
    "- If data is missing, state 'insufficient data' or omit section.\n"
    "- No patient names, dates of birth, or identifying information.\n"
    "- Professional pulmonology / respiratory medicine language.\n"
    "- Cite relevant guidelines: GOLD, GINA, ATS/ERS, Fleischner, CURB-65.\n"
    "- Ready to paste into patient record.\n\n"
)


# ---------------------------------------------------------------------------
# PROMPTS
# ---------------------------------------------------------------------------
def _p_intake(notes, lang, country, style, symptom, smoking, pack_years, occupation):
    return (
        "You are a pulmonologist writing a RESPIRATORY INTAKE ASSESSMENT.\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "Main symptom: " + symptom + "\n"
        "Smoking: " + smoking + " | Pack-years: " + str(pack_years) + "\n"
        "Occupational exposure: " + (occupation or "Not documented") + "\n"
        "Country/system: " + country + "\n\n"
        "Extract from notes and generate:\n\n"
        "1. PRESENTING COMPLAINT + RESPIRATORY HISTORY:\n"
        "   - Symptom characterization\n"
        "   - Duration and progression\n"
        "   - Associated symptoms\n\n"
        "2. RESPIRATORY RISK CLASSIFICATION:\n"
        "   - Smoking risk: low / moderate / high\n"
        "   - Occupational risk: relevant exposures\n"
        "   - Overall pulmonary risk profile\n\n"
        "3. DIFFERENTIAL DIAGNOSIS GROUPS (ranked by probability):\n"
        "   Check: COPD / Asthma / Infection / Malignancy / ILD / PE / Other\n"
        "   - For each group: supporting evidence from notes\n\n"
        "4. ALERT FLAGS:\n"
        "   - SpO2 <92%: hypoxia alert\n"
        "   - Hemoptysis: cancer / TB / PE flag\n"
        "   - Weight loss + smoker: malignancy red flag\n"
        "   - Sudden dyspnea: PE / pneumothorax flag\n\n"
        "5. RECOMMENDED INITIAL WORKUP:\n"
        "   - Immediate tests\n"
        "   - Spirometry indication\n"
        "   - Imaging recommendation\n\n"
        "NOTES:\n" + notes
    )


def _p_spirometry(notes, lang, country, style,
                  fev1_pct, fvc_pct, ratio, dlco_pct):
    # Interpretation logic
    obstruction = ratio < 0.7 if ratio > 0 else None
    restriction = (fvc_pct < 80 and not obstruction) if fvc_pct > 0 else None
    mixed = (ratio < 0.7 and fvc_pct < 80) if (ratio > 0 and fvc_pct > 0) else None

    if fev1_pct > 0 and ratio > 0:
        if ratio < 0.7:
            if fev1_pct >= 80:
                gold_stage = "GOLD I (mild obstruction)"
            elif fev1_pct >= 50:
                gold_stage = "GOLD II (moderate obstruction)"
            elif fev1_pct >= 30:
                gold_stage = "GOLD III (severe obstruction)"
            else:
                gold_stage = "GOLD IV (very severe obstruction)"
        else:
            gold_stage = "No obstruction"
    else:
        gold_stage = "Insufficient data for GOLD staging"

    pattern = "Insufficient data"
    if ratio > 0 and fvc_pct > 0:
        if ratio < 0.7 and fvc_pct >= 80:
            pattern = "OBSTRUCTIVE"
        elif ratio < 0.7 and fvc_pct < 80:
            pattern = "MIXED (obstructive + restrictive)"
        elif ratio >= 0.7 and fvc_pct < 80:
            pattern = "RESTRICTIVE"
        else:
            pattern = "NORMAL"

    return (
        "You are a pulmonologist writing a SPIROMETRY / PULMONARY FUNCTION TEST REPORT.\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "PFT Values entered by physician:\n"
        "FEV1: " + (str(fev1_pct) + "% predicted" if fev1_pct > 0 else "not entered") + "\n"
        "FVC: " + (str(fvc_pct) + "% predicted" if fvc_pct > 0 else "not entered") + "\n"
        "FEV1/FVC ratio: " + (str(round(ratio, 2)) if ratio > 0 else "not entered") + "\n"
        "DLCO: " + (str(dlco_pct) + "% predicted" if dlco_pct > 0 else "not entered") + "\n\n"
        "Pre-calculated interpretation:\n"
        "Pattern: " + pattern + "\n"
        "GOLD stage (if obstructive): " + gold_stage + "\n\n"
        "Additional findings from notes:\n" + (notes or "None") + "\n\n"
        "Write a complete SPIROMETRY INTERPRETATION REPORT:\n\n"
        "1. TECHNICAL QUALITY: acceptability and reproducibility (from notes if mentioned)\n\n"
        "2. RESULTS TABLE:\n"
        "   | Parameter | Value | % Predicted | Interpretation |\n"
        "   Include: FEV1, FVC, FEV1/FVC, DLCO, TLC (if available)\n\n"
        "3. PATTERN INTERPRETATION:\n"
        "   - Normal / Obstructive / Restrictive / Mixed\n"
        "   - Severity grading per ATS/ERS criteria\n\n"
        "4. GOLD STAGING (if obstructive):\n"
        "   " + gold_stage + "\n"
        "   Clinical implications\n\n"
        "5. DLCO INTERPRETATION (if available):\n"
        "   - Normal (>70%) / Mildly reduced (60-70%) / Moderately reduced (40-60%) / Severely reduced (<40%)\n"
        "   - Clinical significance\n\n"
        "6. BRONCHODILATOR RESPONSE (if mentioned in notes):\n"
        "   - Significant if FEV1 or FVC increases >=12% AND >=200mL\n\n"
        "7. CLINICAL CONCLUSION:\n"
        "   - Most likely diagnosis\n"
        "   - Recommended next steps\n"
    )


def _p_copd(notes, lang, country, style, gold_stage, asthma_ctrl, exac_history, pack_years):
    return (
        "You are a pulmonologist writing a COPD / ASTHMA MANAGEMENT REPORT.\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "GOLD Stage: " + gold_stage + "\n"
        "Asthma control: " + asthma_ctrl + "\n"
        "Exacerbation history (last 12 months): " + str(exac_history) + "\n"
        "Pack-years: " + str(pack_years) + "\n\n"
        "Extract from notes and generate:\n\n"
        "SECTION A - IF COPD:\n"
        "1. GOLD ABCD RISK GROUP ASSESSMENT:\n"
        "   - Symptoms (mMRC / CAT score if mentioned)\n"
        "   - Exacerbation risk (low: 0-1 mod exac / high: >=2 mod or >=1 severe)\n"
        "   - Group A / B / E assignment\n\n"
        "2. COPD PHARMACOTHERAPY (per GOLD 2024 guideline):\n"
        "   Group A: SABA or SAMA as needed\n"
        "   Group B: LABA or LAMA (or LABA+LAMA)\n"
        "   Group E (high risk): LABA+LAMA; if eos >300: consider ICS+LABA+LAMA\n"
        "   - Current medications vs guideline-recommended\n"
        "   - Inhaler technique comments if mentioned\n\n"
        "3. OXYGEN THERAPY:\n"
        "   - Indication: PaO2 <55 mmHg or SpO2 <88% at rest\n"
        "   - Target: SpO2 88-92% in COPD (avoid over-oxygenation)\n"
        "   - Flow rate recommendation if data available\n\n"
        "4. EXACERBATION MANAGEMENT (if acute):\n"
        "   - Severity: mild / moderate / severe\n"
        "   - Treatment: SABA, systemic steroids (5 days), antibiotics (if purulent sputum)\n"
        "   - Hospitalization criteria\n\n"
        "5. NON-PHARMACOLOGICAL:\n"
        "   - Pulmonary rehabilitation\n"
        "   - Smoking cessation (if applicable)\n"
        "   - Vaccinations (flu, pneumococcal, COVID)\n\n"
        "SECTION B - IF ASTHMA:\n"
        "1. GINA CONTROL ASSESSMENT:\n"
        "   - Controlled: all criteria met\n"
        "   - Partly controlled: 1-2 criteria\n"
        "   - Uncontrolled: 3-4 criteria\n\n"
        "2. GINA STEP THERAPY (2024):\n"
        "   Step 1-2: low-dose ICS-formoterol (preferred) or SABA\n"
        "   Step 3: low-dose ICS-LABA\n"
        "   Step 4: medium/high ICS-LABA\n"
        "   Step 5: add-on therapy (tiotropium, biologics, low-dose OCS)\n\n"
        "3. TRIGGER MANAGEMENT:\n"
        "   - Allergen, occupational, exercise-induced\n\n"
        "NOTES:\n" + notes
    )


def _p_imaging(notes, lang, country, style):
    return (
        "You are a pulmonologist interpreting THORACIC IMAGING (CT / Chest X-ray).\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "CRITICAL: Extract ONLY findings mentioned in notes. Do NOT invent radiological findings.\n"
        "If no imaging data provided, state 'insufficient data'.\n\n"
        "Interpret and structure the imaging findings:\n\n"
        "1. IMAGING TYPE AND QUALITY:\n"
        "   - Modality (PA CXR / lateral / CT thorax / HRCT / CT-PA / PET-CT)\n"
        "   - Technical quality if mentioned\n\n"
        "2. STRUCTURED FINDINGS:\n"
        "   LUNG FIELDS:\n"
        "   - Parenchyma: clear / consolidation / infiltrate / effusion / pneumothorax\n"
        "   - Pattern: ground-glass / reticular / nodular / cavitary / tree-in-bud\n"
        "   - Location: upper / middle / lower zones, R/L/bilateral\n"
        "   - Air trapping / hyperinflation / emphysema\n"
        "   PLEURA:\n"
        "   - Effusion: size (small/moderate/large), unilateral/bilateral\n"
        "   - Pneumothorax: size estimation\n"
        "   - Pleural thickening / calcification\n"
        "   MEDIASTINUM:\n"
        "   - Widening, lymphadenopathy (stations if CT)\n"
        "   - Cardiac silhouette\n"
        "   NODULES / MASSES:\n"
        "   - Size, location, morphology, density\n"
        "   - Malignant features: speculation, lobulation, pleural traction\n\n"
        "3. IMPRESSION / MOST LIKELY DIAGNOSES:\n"
        "   - Pneumonia pattern (lobar/interstitial/atypical)\n"
        "   - Malignancy suspicion\n"
        "   - Fibrosis / ILD pattern (UIP / NSIP / other)\n"
        "   - Emphysema distribution (centrilobular/panlobular)\n"
        "   - PE findings if CT-PA\n\n"
        "4. FOLLOW-UP RECOMMENDATIONS:\n"
        "   - Based on findings\n"
        "   - Urgency: immediate / 3 months / 6 months / routine\n\n"
        "NOTES (imaging report / findings):\n" + notes
    )


def _p_nodule(notes, lang, country, style,
              size_mm, nodule_type, growth_known, prev_size_mm):
    # Fleischner guideline logic
    fleischner = ""
    if size_mm > 0:
        if size_mm < 6:
            fleischner = "Fleischner 2017: <6mm solid nodule - No routine follow-up needed (low-risk patient)"
        elif size_mm <= 8:
            fleischner = "Fleischner 2017: 6-8mm solid nodule - CT at 6-12 months, then 18-24 months if stable"
        else:
            fleischner = "Fleischner 2017: >8mm solid nodule - CT at 3 months, PET-CT or tissue sampling depending on risk"

        if "ground" in nodule_type.lower() or "glass" in nodule_type.lower() or "milchglas" in nodule_type.lower():
            if size_mm < 6:
                fleischner = "Fleischner 2017: Pure GGO <6mm - No follow-up needed"
            else:
                fleischner = "Fleischner 2017: Pure GGO >=6mm - CT at 6-12 months to confirm persistence, then 2-5 year surveillance"

        if growth_known and prev_size_mm > 0 and size_mm > prev_size_mm:
            vdt = ""
            if prev_size_mm > 0:
                try:
                    ratio = size_mm / prev_size_mm
                    vdt = " | Volume doubling time calculation: suspicious if VDT <400 days"
                except Exception:
                    pass
            fleischner += " | GROWTH DETECTED - upgraded workup recommended" + vdt

    return (
        "You are a pulmonologist / thoracic oncology specialist assessing a LUNG NODULE / PULMONARY MASS.\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "Nodule data entered:\n"
        "Size: " + (str(size_mm) + " mm" if size_mm > 0 else "not entered") + "\n"
        "Type: " + nodule_type + "\n"
        "Known growth: " + ("YES - previous size " + str(prev_size_mm) + " mm" if growth_known and prev_size_mm > 0 else "NO / unknown") + "\n\n"
        "Fleischner 2017 Guideline Assessment:\n"
        + fleischner + "\n\n"
        "Additional clinical context:\n" + (notes or "None") + "\n\n"
        "Generate comprehensive LUNG NODULE ASSESSMENT:\n\n"
        "1. NODULE CHARACTERIZATION:\n"
        "   - Size, type, density, morphology\n"
        "   - Benign features (smooth, calcified, stable) vs malignant features (spiculated, growing)\n\n"
        "2. RISK STRATIFICATION:\n"
        "   - Low risk / Intermediate risk / High risk\n"
        "   - Based on: size, morphology, patient risk factors (smoking, age, family history)\n"
        "   - Malignancy probability estimate (Brock/Mayo model context if applicable)\n\n"
        "3. FLEISCHNER SOCIETY GUIDELINE (2017):\n"
        "   " + fleischner + "\n"
        "   - Apply to this specific nodule\n"
        "   - State clearly: recommended follow-up interval and modality\n\n"
        "4. ADDITIONAL WORKUP:\n"
        "   - PET-CT indication: solid nodule >8mm or high-risk features\n"
        "   - Biopsy/bronchoscopy indication\n"
        "   - MDT/tumor board referral if malignancy suspected\n\n"
        "5. LUNG CANCER STAGING (if malignancy confirmed or highly suspected):\n"
        "   - TNM 8th edition if data available\n"
        "   - Molecular testing needed (EGFR, ALK, KRAS, PD-L1)\n\n"
        "6. DOCUMENTATION FOR FOLLOW-UP:\n"
        "   - Date of detection\n"
        "   - Next imaging date\n"
        "   - Escalation criteria\n"
    )


def _p_infection(notes, lang, country, style, infect_type, curb_score):
    curb_risk = ""
    if curb_score == 0 or curb_score == 1:
        curb_risk = "CURB-65: " + str(curb_score) + " - LOW risk. Outpatient treatment possible."
    elif curb_score == 2:
        curb_risk = "CURB-65: 2 - MODERATE risk. Consider short hospitalization or supervised outpatient."
    elif curb_score >= 3:
        curb_risk = "CURB-65: " + str(curb_score) + " - HIGH risk. Hospitalization required. Consider ICU if 4-5."

    return (
        "You are a pulmonologist / infectious disease specialist writing an INFECTIOUS LUNG DISEASE ASSESSMENT.\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "Infection type: " + infect_type + "\n"
        "CURB-65: " + str(curb_score) + " points - " + curb_risk + "\n\n"
        "Extract from notes and generate:\n\n"
        "1. CLINICAL PRESENTATION:\n"
        "   - Symptoms: fever, cough, sputum, dyspnea, pleurisy\n"
        "   - Duration and progression\n"
        "   - Vital signs (extract from notes)\n\n"
        "2. MICROBIOLOGICAL DATA (if in notes):\n"
        "   - Culture results\n"
        "   - PCR / serology\n"
        "   - Antibiogram if available\n\n"
        "3. IMAGING CORRELATION (if mentioned):\n"
        "   - CXR / CT pattern (lobar / interstitial / cavitary)\n\n"
        "4. CURB-65 ASSESSMENT:\n"
        "   " + curb_risk + "\n"
        "   - Hospitalization decision\n"
        "   - ICU consideration if score >=4\n\n"
        "5. ANTIBIOTIC THERAPY (guideline-based, not automatic prescription):\n"
        "   CAP outpatient: Amoxicillin +/- macrolide (or doxycycline)\n"
        "   CAP inpatient non-severe: Beta-lactam + macrolide\n"
        "   CAP severe / ICU: Beta-lactam + macrolide + fluoroquinolone\n"
        "   HAP/VAP: broad-spectrum per local resistance patterns\n"
        "   TB: 2HRZE / 4HR (refer to specialist)\n"
        "   - Duration of treatment\n"
        "   - Adjust based on microbiological results\n\n"
        "6. TUBERCULOSIS CONSIDERATIONS (if suspected):\n"
        "   - Isolation precautions: YES/NO (airborne isolation)\n"
        "   - Mandatory reporting\n"
        "   - Contact tracing\n"
        "   - IGRA / TST result if available\n\n"
        "7. MONITORING:\n"
        "   - Clinical response criteria (fever defervescence, oxygenation)\n"
        "   - CXR follow-up timing\n"
        "   - Treatment failure criteria\n\n"
        "NOTES:\n" + notes
    )


def _p_icu(notes, lang, country, style, rf_type, spo2, po2, pco2, rr_resp, o2_support, fio2):
    # Respiratory failure classification
    pf_ratio = ""
    if po2 > 0 and fio2 > 0:
        pf = round((po2 / (fio2 / 100)), 1)
        pf_ratio = "P/F ratio: " + str(pf)
        if pf < 100:
            pf_ratio += " -> SEVERE ARDS"
        elif pf < 200:
            pf_ratio += " -> MODERATE ARDS"
        elif pf < 300:
            pf_ratio += " -> MILD ARDS / Type I failure"

    # O2 therapy recommendation based on SpO2
    o2_rec = ""
    if spo2 > 0:
        if spo2 < 88:
            o2_rec = "CRITICAL HYPOXIA - immediate O2 escalation required"
        elif spo2 < 92:
            o2_rec = "Hypoxia - O2 supplementation required, monitor closely"
        elif spo2 > 96 and "COPD" in rf_type:
            o2_rec = "Caution: COPD patient - avoid over-oxygenation, target SpO2 88-92%"

    return (
        "You are a pulmonologist / intensivist writing a RESPIRATORY FAILURE / ICU ASSESSMENT.\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "Respiratory failure type: " + rf_type + "\n"
        "Current O2 support: " + o2_support + "\n\n"
        "Values entered:\n"
        "SpO2: " + (str(spo2) + "%" if spo2 > 0 else "not entered") + "\n"
        "PaO2: " + (str(po2) + " mmHg" if po2 > 0 else "not entered") + "\n"
        "PaCO2: " + (str(pco2) + " mmHg" if pco2 > 0 else "not entered") + "\n"
        "Respiratory rate: " + (str(rr_resp) + "/min" if rr_resp > 0 else "not entered") + "\n"
        "FiO2: " + str(fio2) + "%\n"
        + (pf_ratio + "\n" if pf_ratio else "") +
        (o2_rec + "\n" if o2_rec else "") + "\n"
        "Additional context:\n" + (notes or "None") + "\n\n"
        "Generate RESPIRATORY FAILURE ASSESSMENT:\n\n"
        "1. CLASSIFICATION:\n"
        "   Type I (Hypoxemic): PaO2 <60 mmHg with normal/low PaCO2\n"
        "   Type II (Hypercapnic): PaCO2 >45 mmHg with respiratory acidosis\n"
        "   - Classify this patient based on values\n"
        "   - Severity: mild / moderate / severe\n\n"
        "2. CAUSE / DIFFERENTIAL (from notes):\n"
        "   Type I causes: pneumonia, ARDS, PE, pulmonary edema, pneumothorax\n"
        "   Type II causes: COPD, asthma, NMD, obesity hypoventilation, opioids\n\n"
        "3. OXYGENATION THERAPY PATHWAY:\n"
        "   Step 1: Nasal cannula (1-6 L/min) -> target SpO2 94-98% (88-92% in COPD)\n"
        "   Step 2: Simple/Venturi mask (24-60% FiO2)\n"
        "   Step 3: Non-rebreather mask (60-90% FiO2)\n"
        "   Step 4: HFNC (up to 60 L/min, up to 100% FiO2) - ROX index monitoring\n"
        "   Step 5: NIV / CPAP / BiPAP\n"
        "     - COPD exacerbation: BiPAP (IPAP 14-20 cmH2O, EPAP 4-8 cmH2O)\n"
        "     - Cardiogenic pulmonary edema: CPAP 5-10 cmH2O\n"
        "   Step 6: Endotracheal intubation - ESCALATION CRITERIA\n\n"
        "4. NIV PROTOCOL (if applicable):\n"
        "   - Settings recommendation based on diagnosis\n"
        "   - Monitoring: SpO2, RR, mental status, ABG at 1h\n"
        "   - Failure criteria (need for intubation)\n\n"
        "5. INTUBATION CRITERIA (if present):\n"
        "   - Airway protection failure\n"
        "   - Respiratory arrest / impending\n"
        "   - Failure of NIV\n"
        "   - Hemodynamic instability\n\n"
        "6. ARDS MANAGEMENT (if applicable):\n"
        "   " + (pf_ratio if pf_ratio else "P/F ratio not calculable") + "\n"
        "   - Lung-protective ventilation: TV 6 mL/kg IBW, Pplat <30 cmH2O\n"
        "   - PEEP strategy\n"
        "   - Prone positioning if P/F <150\n\n"
        "7. MONITORING PARAMETERS:\n"
        "   - ABG frequency\n"
        "   - ROX index (if HFNC): SpO2/FiO2 / RR\n"
        "   - Escalation triggers\n"
    )


def _p_sleep(notes, lang, country, style, sb_score, risk_level):
    return (
        "You are a sleep medicine specialist / pulmonologist writing a SLEEP APNEA ASSESSMENT.\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "STOP-BANG Score: " + str(sb_score) + "/8 -> " + risk_level + "\n\n"
        "Additional clinical context:\n" + (notes or "None") + "\n\n"
        "Generate SLEEP MEDICINE ASSESSMENT:\n\n"
        "1. STOP-BANG INTERPRETATION:\n"
        "   0-2: Low risk for OSA\n"
        "   3-4: Intermediate risk\n"
        "   5-8: High risk\n"
        "   This patient: " + risk_level + "\n\n"
        "2. CLINICAL SYMPTOMS ASSESSMENT (from notes):\n"
        "   - Snoring frequency and volume\n"
        "   - Witnessed apneas\n"
        "   - Excessive daytime sleepiness (Epworth scale if mentioned)\n"
        "   - Nocturia, morning headaches, cognitive changes\n\n"
        "3. DIAGNOSTIC RECOMMENDATION:\n"
        "   Low risk: reassure, lifestyle advice\n"
        "   Intermediate risk: home sleep apnea test (HSAT)\n"
        "   High risk: in-lab polysomnography or HSAT + clinical correlation\n"
        "   Urgent PSG: suspected central apnea, COPD overlap, severe hypoxia\n\n"
        "4. AHI CLASSIFICATION (if test result available in notes):\n"
        "   Normal: AHI <5\n"
        "   Mild OSA: AHI 5-14\n"
        "   Moderate OSA: AHI 15-29\n"
        "   Severe OSA: AHI >=30\n\n"
        "5. TREATMENT RECOMMENDATION:\n"
        "   CPAP: first-line for moderate-severe OSA; also mild if symptomatic\n"
        "   - Pressure range (auto-CPAP typical range: 4-20 cmH2O)\n"
        "   - BiPAP if CPAP intolerant or obesity hypoventilation\n"
        "   - Mandibular advancement device: mild-moderate OSA, CPAP intolerant\n"
        "   - Positional therapy: position-dependent OSA\n\n"
        "6. CARDIOVASCULAR RISK:\n"
        "   - OSA association with: HTN, AF, CAD, stroke\n"
        "   - CPAP cardiovascular benefit context\n\n"
        "7. FOLLOW-UP:\n"
        "   - CPAP compliance check: 1-3 months\n"
        "   - Epworth score reassessment\n"
        "   - Driving restrictions if severe OSA (local regulations)\n"
    )


def _p_followup(notes, lang, country, style):
    return (
        "You are a pulmonologist writing a FOLLOW-UP AND MONITORING REPORT.\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "Extract from notes and generate a comprehensive follow-up assessment:\n\n"
        "1. DISEASE PROGRESSION TRACKING:\n"
        "   COPD: FEV1 trend, exacerbation frequency, hospitalization history\n"
        "   Asthma: symptom control, exacerbations, SABA use\n"
        "   ILD/Fibrosis: FVC trend, DLCO trend, 6MWT distance\n"
        "   Lung cancer: response to treatment, tumor markers, imaging interval\n\n"
        "2. SPIROMETRY TREND (if serial values in notes):\n"
        "   - FEV1 change from baseline\n"
        "   - Rate of decline (FEV1 ml/year normal: ~30 ml/year; accelerated in COPD)\n\n"
        "3. EXACERBATION HISTORY:\n"
        "   - Number and severity in last 12 months\n"
        "   - Triggers identified\n"
        "   - Prevention strategy\n\n"
        "4. CURRENT TREATMENT ASSESSMENT:\n"
        "   - Medication adherence and technique\n"
        "   - Response to current therapy\n"
        "   - Side effects reported\n"
        "   - Step up / step down recommendation\n\n"
        "5. NODULE FOLLOW-UP (if applicable):\n"
        "   - Size comparison to previous CT\n"
        "   - Volume doubling time calculation if growth\n"
        "   - Next imaging recommendation\n\n"
        "6. OXYGEN THERAPY REVIEW (if applicable):\n"
        "   - Hours per day compliance\n"
        "   - Reassessment: still needed?\n\n"
        "7. QUALITY OF LIFE / FUNCTIONAL STATUS:\n"
        "   - Exercise tolerance (MRC dyspnea scale)\n"
        "   - CAT score if mentioned\n"
        "   - Activities of daily living\n\n"
        "8. NEXT FOLLOW-UP PLAN:\n"
        "   - Timing of next visit\n"
        "   - Tests needed\n"
        "   - Referrals\n\n"
        "NOTES:\n" + notes
    )


def _p_discharge(notes, lang, country, style, scenario):
    return (
        "You are a pulmonologist writing a PULMONARY DISCHARGE LETTER.\n" +
        RULES.format(lang=LANG_NAME.get(lang,"German"), style=style) +
        "Discharge scenario: " + scenario + "\n"
        "Country/system: " + country + "\n\n"
        "Extract from notes and write a complete discharge letter:\n\n"
        "1. HEADER: clinic/department, physician, patient age/sex only (NO name/DOB)\n\n"
        "2. MAIN DIAGNOSES (ICD-10 codes if country uses them)\n\n"
        "3. REASON FOR HOSPITALIZATION\n\n"
        "4. KEY FINDINGS:\n"
        "   - Spirometry results\n"
        "   - Imaging\n"
        "   - Microbiology\n"
        "   - O2 requirements\n"
        "   - Sleep study results if applicable\n\n"
        "5. HOSPITAL COURSE:\n"
        "   - Treatment given\n"
        "   - Response to therapy\n"
        "   - Complications\n\n"
        "6. DISCHARGE CONDITION:\n"
        "   - SpO2 at discharge\n"
        "   - O2 requirement\n"
        "   - Functional status\n\n"
        "7. MEDICATION RECONCILIATION:\n"
        "   | Medication | Admission | Discharge | Change reason |\n"
        "   Include: inhalers (with device type), O2 prescription, antibiotics\n\n"
        "8. SPECIFIC DISCHARGE INSTRUCTIONS:\n"
        "   COPD: inhaler technique, action plan for exacerbation\n"
        "   Asthma: written asthma action plan\n"
        "   Pneumonia: completion of antibiotics, CXR follow-up\n"
        "   Cancer: next appointment, chemotherapy/immunotherapy plan\n"
        "   ICU: weaning parameters, home O2 or NIV\n\n"
        "9. FOLLOW-UP PLAN:\n"
        "   - Pulmonology outpatient: when\n"
        "   - Repeat spirometry\n"
        "   - CT thorax follow-up\n"
        "   - GP visit\n\n"
        "10. RED FLAGS for re-admission:\n"
        "    - SpO2 drop below threshold\n"
        "    - Increased dyspnea\n"
        "    - Hemoptysis\n\n"
        "11. LEGAL FOOTER\n\n"
        "NOTES:\n" + notes
    )


# ---------------------------------------------------------------------------
# TAB RENDERERS
# ---------------------------------------------------------------------------
def _tab_intake(lang, client, country, style):
    st.markdown("#### " + {
        "de":"Pneumologische Aufnahme / Risikoeinstufung",
        "en":"Pulmonary intake / risk classification",
        "ro":"Admitere pneumologica / clasificare risc",
        "es":"Ingreso neumologico / clasificacion riesgo",
        "ar":"  /  "
    }.get(lang,"Pulmonary intake"))

    c1, c2, c3 = st.columns(3)
    with c1:
        s_opts = tx("symp_opts", lang)
        st.session_state.pn_symptom = st.selectbox(
            tx("main_symptom", lang), range(len(s_opts)),
            format_func=lambda i: s_opts[i],
            index=st.session_state.pn_symptom, key="pn_sel_symp"
        )
    with c2:
        sm_opts = tx("smoking_opts", lang)
        st.session_state.pn_smoking = st.selectbox(
            tx("smoking", lang), range(len(sm_opts)),
            format_func=lambda i: sm_opts[i],
            index=st.session_state.pn_smoking, key="pn_sel_smoke"
        )
    with c3:
        st.session_state.pn_pack_years = st.number_input(
            tx("pack_years", lang), 0.0, 200.0,
            value=st.session_state.pn_pack_years, step=1.0, key="pn_py"
        )

    if st.session_state.pn_pack_years >= 20 and st.session_state.pn_smoking != 0:
        st.warning({
            "de":"Erhoehtes Risiko: >=20 Packungsjahre - Malignomabklaerung erwaegen",
            "en":"Elevated risk: >=20 pack-years - consider malignancy workup",
            "ro":"Risc crescut: >=20 pachete-ani - considerati bilant oncologic",
            "es":"Riesgo elevado: >=20 paquetes-ano - considere estudio oncologico",
            "ar":" : 20+   -   "
        }.get(lang,"High smoking risk"))

    st.session_state.pn_occupation = st.text_input(
        tx("occupation", lang),
        value=st.session_state.pn_occupation, key="pn_occ"
    )

    symptom = s_opts[st.session_state.pn_symptom]
    smoking = sm_opts[st.session_state.pn_smoking]

    st.info({
        "de":"KI klassifiziert Risiko, bewertet Verdachtsgruppen (COPD/Asthma/Infektion/Tumor/ILD), erkennt Red Flags.",
        "en":"AI classifies risk, assesses suspected groups (COPD/Asthma/Infection/Cancer/ILD), detects red flags.",
        "ro":"AI clasifica riscul, evalueaza grupe suspecte (BPOC/Astm/Infectie/Tumor/ILD), detecteaza red flags.",
        "es":"La IA clasifica el riesgo, evalua grupos sospechosos (EPOC/Asma/Infeccion/Cancer/EPI), detecta red flags.",
        "ar":"       (COPD////ILD)   ."
    }.get(lang,"AI classifies risk and detects red flags."))

    notes = _notes_and_audio(lang, "intake", client)

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_intake"):
        if not notes.strip():
            st.warning(tx("fill_warn", lang))
        else:
            with st.spinner(tx("generating", lang)):
                try:
                    st.session_state.pn_res_intake = _call_ai(
                        _p_intake(notes, lang, country, style,
                                  symptom, smoking,
                                  st.session_state.pn_pack_years,
                                  st.session_state.pn_occupation),
                        client, max_tokens=2000
                    )
                except Exception as e:
                    st.error(str(e))

    _output_block("pn_res_intake", {
        "de":"Pneumologische Aufnahme","en":"Pulmonary Intake Assessment",
        "ro":"Evaluare Admitere Pneumologica","es":"Evaluacion Ingreso Neumologico","ar":"  "
    }.get(lang,"Intake"), lang)


def _tab_spirometry(lang, client, country, style):
    st.markdown("#### " + {
        "de":"Spirometrie / Lungenfunktion (PFT)",
        "en":"Spirometry / Pulmonary Function Testing (PFT)",
        "ro":"Spirometrie / Testare Functionala Pulmonara",
        "es":"Espirometria / Funcion Pulmonar (PFT)",
        "ar":"  /  "
    }.get(lang,"Spirometry / PFT"))

    st.info({
        "de":"KI interpretiert: obstruktiv/restriktiv/gemischt/normal, GOLD-Staging, DLCO, Bronchodilatationstest.",
        "en":"AI interprets: obstructive/restrictive/mixed/normal, GOLD staging, DLCO, bronchodilator response.",
        "ro":"AI interpreteaza: obstructiv/restrictiv/mixt/normal, stadializare GOLD, DLCO, raspuns bronhodilatator.",
        "es":"La IA interpreta: obstructivo/restrictivo/mixto/normal, estadificacion GOLD, DLCO, respuesta broncodilatadora.",
        "ar":"  : ///  GOLD DLCO   ."
    }.get(lang,"AI interprets all spirometry data."))

    c1, c2, c3 = st.columns(3)
    with c1:
        st.session_state.pn_fev1_pct = st.number_input(
            "FEV1 (% predicted)", 0.0, 200.0,
            value=st.session_state.pn_fev1_pct, step=1.0, key="pn_fev1p"
        )
        st.session_state.pn_fvc_pct = st.number_input(
            "FVC (% predicted)", 0.0, 200.0,
            value=st.session_state.pn_fvc_pct, step=1.0, key="pn_fvcp"
        )
    with c2:
        st.session_state.pn_fev1_fvc = st.number_input(
            "FEV1/FVC ratio", 0.0, 1.5,
            value=st.session_state.pn_fev1_fvc, step=0.01, key="pn_ratio"
        )
        st.session_state.pn_dlco_pct = st.number_input(
            "DLCO (% predicted)", 0.0, 200.0,
            value=st.session_state.pn_dlco_pct, step=1.0, key="pn_dlco"
        )
    with c3:
        # Live interpretation
        ratio = st.session_state.pn_fev1_fvc
        fev1p = st.session_state.pn_fev1_pct
        fvcp = st.session_state.pn_fvc_pct
        if ratio > 0 and fev1p > 0:
            if ratio < 0.7:
                if fev1p >= 80:
                    st.error("GOLD I - Mild obstruction")
                elif fev1p >= 50:
                    st.warning("GOLD II - Moderate obstruction")
                elif fev1p >= 30:
                    st.error("GOLD III - Severe obstruction")
                else:
                    st.error("GOLD IV - Very severe")
            elif fvcp < 80:
                st.warning("Restrictive pattern")
            else:
                st.success("Normal spirometry")
        else:
            st.caption("Enter values for live interpretation")

    notes = _notes_and_audio(lang, "spiro", client, height=120,
        placeholder_extra=" (z.B. Bronchodilatationstest, 6MWT, TLC)")

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_spiro"):
        with st.spinner(tx("generating", lang)):
            try:
                st.session_state.pn_res_spiro = _call_ai(
                    _p_spirometry(notes, lang, country, style,
                                  st.session_state.pn_fev1_pct,
                                  st.session_state.pn_fvc_pct,
                                  st.session_state.pn_fev1_fvc,
                                  st.session_state.pn_dlco_pct),
                    client, max_tokens=2000
                )
            except Exception as e:
                st.error(str(e))

    _output_block("pn_res_spiro", {
        "de":"Spirometrie-Bericht","en":"Spirometry Report",
        "ro":"Raport Spirometrie","es":"Informe Espirometria","ar":"  "
    }.get(lang,"Spirometry Report"), lang)


def _tab_copd(lang, client, country, style):
    st.markdown("#### " + {
        "de":"COPD / Asthma Management Engine",
        "en":"COPD / Asthma Management Engine",
        "ro":"Motor Management BPOC / Astm",
        "es":"Motor Manejo EPOC / Asma",
        "ar":"  COPD / "
    }.get(lang,"COPD / Asthma"))

    mode_label = {
        "de":["COPD Modus","Asthma Modus"],
        "en":["COPD Mode","Asthma Mode"],
        "ro":["Mod BPOC","Mod Astm"],
        "es":["Modo EPOC","Modo Asma"],
        "ar":[" COPD"," "]
    }.get(lang,["COPD Mode","Asthma Mode"])
    mode = st.radio("", mode_label, horizontal=True, key="pn_copd_mode",
                    label_visibility="collapsed")

    if mode == mode_label[0]:
        c1, c2 = st.columns(2)
        with c1:
            gold_opts = tx("gold_opts", lang)
            st.session_state.pn_gold = st.selectbox(
                tx("copd_stage", lang), range(len(gold_opts)),
                format_func=lambda i: gold_opts[i],
                index=st.session_state.pn_gold, key="pn_sel_gold"
            )
        with c2:
            st.session_state.pn_exac_history = st.number_input(
                {"de":"Exazerbationen (letzte 12 Monate)","en":"Exacerbations (last 12 months)","ro":"Exacerbari (ultimele 12 luni)","es":"Exacerbaciones (ultimos 12 meses)","ar":" ( 12 )"}.get(lang,"Exacerbations"),
                0, 20, value=st.session_state.pn_exac_history, key="pn_exac"
            )
        gold_stage = gold_opts[st.session_state.pn_gold]
        asthma_ctrl = "N/A"
    else:
        c1, = st.columns([1])
        asthma_opts = tx("asthma_opts", lang)
        st.session_state.pn_asthma_ctrl = st.selectbox(
            tx("asthma_ctrl", lang), range(len(asthma_opts)),
            format_func=lambda i: asthma_opts[i],
            index=st.session_state.pn_asthma_ctrl, key="pn_sel_asthma"
        )
        gold_stage = "N/A (Asthma)"
        asthma_ctrl = asthma_opts[st.session_state.pn_asthma_ctrl]

    notes = _notes_and_audio(lang, "copd", client)

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_copd"):
        if not notes.strip():
            st.warning(tx("fill_warn", lang))
        else:
            with st.spinner(tx("generating", lang)):
                try:
                    st.session_state.pn_res_copd = _call_ai(
                        _p_copd(notes, lang, country, style,
                                gold_stage, asthma_ctrl,
                                st.session_state.pn_exac_history,
                                st.session_state.pn_pack_years),
                        client, max_tokens=2500
                    )
                except Exception as e:
                    st.error(str(e))

    _output_block("pn_res_copd", {
        "de":"COPD / Asthma Management Report","en":"COPD / Asthma Management Report",
        "ro":"Raport Management BPOC / Astm","es":"Informe Manejo EPOC / Asma","ar":"  COPD / "
    }.get(lang,"COPD/Asthma Report"), lang)


def _tab_imaging(lang, client, country, style):
    st.markdown("#### " + {
        "de":"Bildgebungs-Interpretation (CT Thorax / Roentgen)",
        "en":"Imaging interpretation (CT Thorax / X-ray)",
        "ro":"Interpretare imagistica (CT Toracic / Radiografie)",
        "es":"Interpretacion imagen (TC Toracico / Radiografia)",
        "ar":"  (CT  / )"
    }.get(lang,"Imaging"))

    st.warning({
        "de":"Nur Befunde aus den Notizen werden verwendet. KI erfindet KEINE radiologischen Befunde.",
        "en":"Only findings from notes are used. AI does NOT invent radiological findings.",
        "ro":"Doar constatari din note. AI NU inventa rezultate radiologice.",
        "es":"Solo hallazgos de las notas. La IA NO inventa hallazgos radiologicos.",
        "ar":"   .      ."
    }.get(lang,"Findings from notes only."))

    st.info({
        "de":"KI extrahiert: Lungenfelder, Pleura, Mediastinum, Noduli, Befundmuster - strukturiert nach WHO/Fleischner Standards.",
        "en":"AI extracts: lung fields, pleura, mediastinum, nodules, patterns - structured per WHO/Fleischner standards.",
        "ro":"AI extrage: campuri pulmonare, pleura, mediastin, noduli, tipare - structurat per standarde WHO/Fleischner.",
        "es":"La IA extrae: campos pulmonares, pleura, mediastino, nodulos, patrones - estructurado segun estandares WHO/Fleischner.",
        "ar":"  :       -   Fleischner."
    }.get(lang,"AI structures imaging findings."))

    notes = _notes_and_audio(lang, "imaging", client, height=200,
        placeholder_extra=" (CT-Befund, Roentgen-Beschreibung, Radiologen-Bericht)")

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_img"):
        if not notes.strip():
            st.warning(tx("fill_warn", lang))
        else:
            with st.spinner(tx("generating", lang)):
                try:
                    st.session_state.pn_res_imaging = _call_ai(
                        _p_imaging(notes, lang, country, style),
                        client, max_tokens=2000
                    )
                except Exception as e:
                    st.error(str(e))

    _output_block("pn_res_imaging", {
        "de":"Bildgebungs-Bericht","en":"Imaging Report",
        "ro":"Raport Imagistica","es":"Informe Imagen","ar":" "
    }.get(lang,"Imaging Report"), lang)


def _tab_nodule(lang, client, country, style):
    st.markdown("#### " + {
        "de":"Lungenrundherd / Lungenkarzinom Engine",
        "en":"Lung Nodule / Lung Cancer Engine",
        "ro":"Motor Nodul Pulmonar / Cancer Pulmonar",
        "es":"Motor Nodulo / Cancer Pulmonar",
        "ar":"   /  "
    }.get(lang,"Nodule / Cancer"))

    c1, c2, c3 = st.columns(3)
    with c1:
        st.session_state.pn_nodule_size = st.number_input(
            tx("nodule_size", lang), 0.0, 150.0,
            value=st.session_state.pn_nodule_size, step=0.5, key="pn_nod_sz"
        )
    with c2:
        n_opts = tx("nodule_opts", lang)
        st.session_state.pn_nodule_type = st.selectbox(
            tx("nodule_type", lang), range(len(n_opts)),
            format_func=lambda i: n_opts[i],
            index=st.session_state.pn_nodule_type, key="pn_nod_ty"
        )
    with c3:
        st.session_state.pn_nodule_growth = st.checkbox(
            tx("nodule_growth", lang),
            value=st.session_state.pn_nodule_growth, key="pn_nod_gr"
        )
        if st.session_state.pn_nodule_growth:
            st.session_state.pn_nodule_prev_size = st.number_input(
                {"de":"Vorgroesse (mm)","en":"Previous size (mm)","ro":"Marime anterioara (mm)","es":"Tamano previo (mm)","ar":"  ()"}.get(lang,"Previous size"),
                0.0, 150.0, value=st.session_state.pn_nodule_prev_size, step=0.5, key="pn_nod_prev"
            )

    size = st.session_state.pn_nodule_size
    nodule_type = n_opts[st.session_state.pn_nodule_type]

    # Live Fleischner guidance
    if size > 0:
        if size < 6:
            st.success({
                "de":"Fleischner 2017: <6mm - Kein Routine-Follow-up (Niedrigrisikopatient)",
                "en":"Fleischner 2017: <6mm - No routine follow-up (low-risk patient)",
                "ro":"Fleischner 2017: <6mm - Fara urmarire de rutina",
                "es":"Fleischner 2017: <6mm - Sin seguimiento rutinario",
                "ar":"Fleischner 2017:   6  -   "
            }.get(lang,"Fleischner: No follow-up"))
        elif size <= 8:
            st.warning({
                "de":"Fleischner 2017: 6-8mm - CT Follow-up 6-12 Monate, dann 18-24 Monate",
                "en":"Fleischner 2017: 6-8mm - CT follow-up 6-12 months, then 18-24 months",
                "ro":"Fleischner 2017: 6-8mm - CT la 6-12 luni, apoi 18-24 luni",
                "es":"Fleischner 2017: 6-8mm - TC a los 6-12 meses, luego 18-24 meses",
                "ar":"Fleischner 2017: 6-8  - CT  6-12   18-24 "
            }.get(lang,"Fleischner: CT follow-up"))
        else:
            st.error({
                "de":"Fleischner 2017: >8mm - CT 3 Monate, PET-CT / Biopsie je nach Risiko",
                "en":"Fleischner 2017: >8mm - CT at 3 months, PET-CT or biopsy depending on risk",
                "ro":"Fleischner 2017: >8mm - CT la 3 luni, PET-CT sau biopsie",
                "es":"Fleischner 2017: >8mm - TC a los 3 meses, PET-TC o biopsia",
                "ar":"Fleischner 2017:   8  - CT  3  PET-CT  "
            }.get(lang,"Fleischner: Urgent workup"))

    notes = _notes_and_audio(lang, "nodule", client)

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_nod"):
        with st.spinner(tx("generating", lang)):
            try:
                st.session_state.pn_res_nodule = _call_ai(
                    _p_nodule(notes, lang, country, style,
                              size, nodule_type,
                              st.session_state.pn_nodule_growth,
                              st.session_state.pn_nodule_prev_size),
                    client, max_tokens=2000
                )
            except Exception as e:
                st.error(str(e))

    _output_block("pn_res_nodule", {
        "de":"Lungenrundherd-Assessment","en":"Lung Nodule Assessment",
        "ro":"Evaluare Nodul Pulmonar","es":"Evaluacion Nodulo Pulmonar","ar":"  "
    }.get(lang,"Nodule Assessment"), lang)


def _tab_infection(lang, client, country, style):
    st.markdown("#### " + {
        "de":"Lungeninfektion / Pneumonie Engine",
        "en":"Lung Infection / Pneumonia Engine",
        "ro":"Motor Infectie Pulmonara / Pneumonie",
        "es":"Motor Infeccion Pulmonar / Neumonia",
        "ar":"   /  "
    }.get(lang,"Lung Infection"))

    c1, c2 = st.columns(2)
    with c1:
        i_opts = tx("infect_opts", lang)
        st.session_state.pn_infect_type = st.selectbox(
            tx("infect_type", lang), range(len(i_opts)),
            format_func=lambda i: i_opts[i],
            index=st.session_state.pn_infect_type, key="pn_sel_inf"
        )

    # CURB-65
    st.markdown("**CURB-65**")
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        st.session_state.pn_curb_conf = st.checkbox(
            "C - Confusion", value=st.session_state.pn_curb_conf, key="pn_curb_c"
        )
    with c2:
        st.session_state.pn_curb_urea = st.checkbox(
            "U - Urea >7", value=st.session_state.pn_curb_urea, key="pn_curb_u"
        )
    with c3:
        st.session_state.pn_curb_rr = st.checkbox(
            "R - RR >=30", value=st.session_state.pn_curb_rr, key="pn_curb_r"
        )
    with c4:
        st.session_state.pn_curb_bp = st.checkbox(
            "B - BP low", value=st.session_state.pn_curb_bp, key="pn_curb_b"
        )
    with c5:
        st.session_state.pn_curb_age = st.checkbox(
            "65 - Age >=65", value=st.session_state.pn_curb_age, key="pn_curb_65"
        )

    curb = sum([
        st.session_state.pn_curb_conf, st.session_state.pn_curb_urea,
        st.session_state.pn_curb_rr, st.session_state.pn_curb_bp,
        st.session_state.pn_curb_age
    ])
    if curb <= 1:
        st.success("CURB-65 = " + str(curb) + " -> Low risk - outpatient possible")
    elif curb == 2:
        st.warning("CURB-65 = 2 -> Moderate - consider admission")
    else:
        st.error("CURB-65 = " + str(curb) + " -> High risk - hospitalize" + (" / ICU" if curb >= 4 else ""))

    notes = _notes_and_audio(lang, "infection", client)

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_inf"):
        if not notes.strip():
            st.warning(tx("fill_warn", lang))
        else:
            with st.spinner(tx("generating", lang)):
                try:
                    st.session_state.pn_res_infection = _call_ai(
                        _p_infection(notes, lang, country, style,
                                     i_opts[st.session_state.pn_infect_type], curb),
                        client, max_tokens=2000
                    )
                except Exception as e:
                    st.error(str(e))

    _output_block("pn_res_infection", {
        "de":"Lungeninfektion-Assessment","en":"Pulmonary Infection Assessment",
        "ro":"Evaluare Infectie Pulmonara","es":"Evaluacion Infeccion Pulmonar","ar":"  "
    }.get(lang,"Infection Assessment"), lang)


def _tab_icu(lang, client, country, style):
    st.markdown("#### " + {
        "de":"Atemversagen / ICU / Beatmung",
        "en":"Respiratory Failure / ICU / Ventilation",
        "ro":"Insuficienta Respiratorie / ATI / Ventilatie",
        "es":"Insuficiencia Respiratoria / UCI / Ventilacion",
        "ar":"  / ICU / "
    }.get(lang,"Resp. Failure / ICU"))

    # Critical SpO2 alert
    spo2_val = st.session_state.pn_spo2
    if spo2_val < 88:
        st.error({
            "de":"KRITISCHE HYPOXAEMIE - SpO2 <88% - sofortige Intervention erforderlich",
            "en":"CRITICAL HYPOXEMIA - SpO2 <88% - immediate intervention required",
            "ro":"HIPOXEMIE CRITICA - SpO2 <88% - interventie imediata necesara",
            "es":"HIPOXEMIA CRITICA - SpO2 <88% - intervencion inmediata requerida",
            "ar":"   - SpO2   88% -   "
        }.get(lang,"CRITICAL HYPOXEMIA"))

    c1, c2, c3 = st.columns(3)
    with c1:
        rf_opts = tx("rf_opts", lang)
        st.session_state.pn_rf_type = st.selectbox(
            tx("resp_fail_type", lang), range(len(rf_opts)),
            format_func=lambda i: rf_opts[i],
            index=st.session_state.pn_rf_type, key="pn_sel_rf"
        )
        o2_opts = tx("o2_opts", lang)
        st.session_state.pn_o2_support = st.selectbox(
            tx("o2_support", lang), range(len(o2_opts)),
            format_func=lambda i: o2_opts[i],
            index=st.session_state.pn_o2_support, key="pn_sel_o2"
        )
    with c2:
        st.session_state.pn_spo2 = st.number_input(
            "SpO2 (%)", 50, 100, value=st.session_state.pn_spo2, key="pn_spo2"
        )
        st.session_state.pn_rr_resp = st.number_input(
            {"de":"Atemfrequenz (/min)","en":"Respiratory rate (/min)","ro":"FR (/min)","es":"FR (/min)","ar":" "}.get(lang,"RR /min"),
            4, 60, value=st.session_state.pn_rr_resp, key="pn_rr"
        )
    with c3:
        st.session_state.pn_po2 = st.number_input(
            "PaO2 (mmHg)", 0.0, 700.0, value=st.session_state.pn_po2, step=1.0, key="pn_po2"
        )
        st.session_state.pn_pco2 = st.number_input(
            "PaCO2 (mmHg)", 0.0, 120.0, value=st.session_state.pn_pco2, step=1.0, key="pn_pco2"
        )
        st.session_state.pn_fio2 = st.number_input(
            "FiO2 (%)", 21, 100, value=st.session_state.pn_fio2, key="pn_fio2"
        )

    # P/F ratio live
    if st.session_state.pn_po2 > 0 and st.session_state.pn_fio2 > 0:
        pf = round(st.session_state.pn_po2 / (st.session_state.pn_fio2 / 100), 1)
        if pf < 100:
            st.error("P/F ratio = " + str(pf) + " -> SEVERE ARDS")
        elif pf < 200:
            st.error("P/F ratio = " + str(pf) + " -> MODERATE ARDS")
        elif pf < 300:
            st.warning("P/F ratio = " + str(pf) + " -> MILD ARDS / Type I failure")
        else:
            st.success("P/F ratio = " + str(pf) + " -> Normal oxygenation")

    rf_type = rf_opts[st.session_state.pn_rf_type]
    o2_support = o2_opts[st.session_state.pn_o2_support]
    notes = _notes_and_audio(lang, "icu", client, height=140)

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_icu"):
        with st.spinner(tx("generating", lang)):
            try:
                st.session_state.pn_res_icu = _call_ai(
                    _p_icu(notes, lang, country, style,
                           rf_type, st.session_state.pn_spo2,
                           st.session_state.pn_po2, st.session_state.pn_pco2,
                           st.session_state.pn_rr_resp, o2_support,
                           st.session_state.pn_fio2),
                    client, max_tokens=2500
                )
            except Exception as e:
                st.error(str(e))

    _output_block("pn_res_icu", {
        "de":"Atemversagen / ICU Bericht","en":"Respiratory Failure / ICU Report",
        "ro":"Raport Insuficienta Respiratorie / ATI","es":"Informe Insuficiencia Respiratoria / UCI","ar":"   / ICU"
    }.get(lang,"Resp. Failure Report"), lang)


def _tab_sleep(lang, client, country, style):
    st.markdown("#### " + {
        "de":"Schlafmedizin / Schlafapnoe",
        "en":"Sleep Medicine / Sleep Apnea",
        "ro":"Medicina Somnului / Apnee in Somn",
        "es":"Medicina del Sueno / Apnea del Sueno",
        "ar":"  /    "
    }.get(lang,"Sleep Medicine"))

    st.markdown("**" + tx("stopbang_label", lang) + "**")

    sb_items = [
        ("pn_sb_snore", {"de":"S - Schnarchen laut?","en":"S - Snoring loudly?","ro":"S - Sforait tare?","es":"S - Ronca fuerte?","ar":"S -  "}),
        ("pn_sb_tired", {"de":"T - Muedigkeit / Erschoepfung am Tag?","en":"T - Tired during the day?","ro":"T - Oboseala ziua?","es":"T - Cansancio diurno?","ar":"T -   "}),
        ("pn_sb_obs",   {"de":"O - Beobachtete Atempausen?","en":"O - Observed apnea?","ro":"O - Apnee observata?","es":"O - Apnea observada?","ar":"O -   "}),
        ("pn_sb_bp",    {"de":"P - Bluthochdruck bekannt?","en":"P - High blood pressure?","ro":"P - HTA cunoscuta?","es":"P - HTA conocida?","ar":"P -   "}),
        ("pn_sb_bmi",   {"de":"B - BMI >35?","en":"B - BMI >35?","ro":"B - IMC >35?","es":"B - IMC >35?","ar":"B -   >35"}),
        ("pn_sb_age",   {"de":"A - Alter >50?","en":"A - Age >50?","ro":"A - Varsta >50?","es":"A - Edad >50?","ar":"A -    50"}),
        ("pn_sb_neck",  {"de":"N - Halsumfang >40cm?","en":"N - Neck >40cm?","ro":"N - Circumferinta gat >40cm?","es":"N - Cuello >40cm?","ar":"N -   >40 "}),
        ("pn_sb_male",  {"de":"G - Maennliches Geschlecht?","en":"G - Male gender?","ro":"G - Sex masculin?","es":"G - Sexo masculino?","ar":"G - "}),
    ]

    c1, c2 = st.columns(2)
    sb_total = 0
    for i, (key, label_dict) in enumerate(sb_items):
        with [c1, c2][i % 2]:
            st.session_state[key] = st.checkbox(
                label_dict.get(lang, label_dict.get("en", key)),
                value=st.session_state[key], key="sb_" + key
            )
        if st.session_state[key]:
            sb_total += 1

    if sb_total <= 2:
        risk_level = "Low risk for OSA"
        st.success("STOP-BANG = " + str(sb_total) + "/8 -> " + risk_level)
    elif sb_total <= 4:
        risk_level = "Intermediate risk for OSA"
        st.warning("STOP-BANG = " + str(sb_total) + "/8 -> " + risk_level)
    else:
        risk_level = "High risk for OSA"
        st.error("STOP-BANG = " + str(sb_total) + "/8 -> " + risk_level)

    notes = _notes_and_audio(lang, "sleep", client, height=140,
        placeholder_extra=" (AHI-Wert, PSG-Ergebnis, CPAP-Druck, Epworth-Score)")

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_sleep"):
        with st.spinner(tx("generating", lang)):
            try:
                st.session_state.pn_res_sleep = _call_ai(
                    _p_sleep(notes, lang, country, style, sb_total, risk_level),
                    client, max_tokens=1800
                )
            except Exception as e:
                st.error(str(e))

    _output_block("pn_res_sleep", {
        "de":"Schlafapnoe-Assessment","en":"Sleep Apnea Assessment",
        "ro":"Evaluare Apnee in Somn","es":"Evaluacion Apnea del Sueno","ar":"    "
    }.get(lang,"Sleep Apnea Assessment"), lang)


def _tab_followup(lang, client, country, style):
    st.markdown("#### " + {
        "de":"Verlaufs-Monitoring / Follow-up",
        "en":"Progress Monitoring / Follow-up",
        "ro":"Monitorizare Evolutie / Urmarire",
        "es":"Seguimiento / Monitoreo",
        "ar":"  / "
    }.get(lang,"Follow-up"))

    st.info({
        "de":"KI analysiert: FEV1-Trend, Exazerbationsgeschichte, Tumor-Follow-up, O2-Bedarf, Lebensqualitaet, naechste Schritte.",
        "en":"AI analyzes: FEV1 trend, exacerbation history, tumor follow-up, O2 needs, quality of life, next steps.",
        "ro":"AI analizeaza: tendinta FEV1, istoricul exacerbarilor, urmarire tumor, necesar O2, calitatea vietii.",
        "es":"La IA analiza: tendencia FEV1, historia exacerbaciones, seguimiento tumor, necesidad O2, calidad de vida.",
        "ar":"  :  FEV1      O2  ."
    }.get(lang,"AI analyzes all follow-up data."))

    notes = _notes_and_audio(lang, "followup", client, height=200,
        placeholder_extra=" (FEV1-Verlauf, Exazerbationen, CT-Vergleich, CAT-Score, 6MWT)")

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_fu"):
        if not notes.strip():
            st.warning(tx("fill_warn", lang))
        else:
            with st.spinner(tx("generating", lang)):
                try:
                    st.session_state.pn_res_followup = _call_ai(
                        _p_followup(notes, lang, country, style),
                        client, max_tokens=2000
                    )
                except Exception as e:
                    st.error(str(e))

    _output_block("pn_res_followup", {
        "de":"Verlaufsbericht Pneumologie","en":"Pulmonary Follow-up Report",
        "ro":"Raport Evolutie Pneumologie","es":"Informe Seguimiento Neumologia","ar":"  "
    }.get(lang,"Follow-up Report"), lang)


def _tab_discharge(lang, client, country, style):
    st.markdown("#### " + {
        "de":"Pneumologischer Entlassungsbrief",
        "en":"Pulmonary Discharge Letter",
        "ro":"Scrisoare Externare Pneumologie",
        "es":"Informe Alta Neumologia",
        "ar":"  "
    }.get(lang,"Discharge"))

    dc_opts = tx("dc_opts", lang)
    st.session_state.pn_dc_scenario = st.selectbox(
        tx("dc_scenario", lang), range(len(dc_opts)),
        format_func=lambda i: dc_opts[i],
        index=st.session_state.pn_dc_scenario, key="pn_sel_dc"
    )
    scenario = dc_opts[st.session_state.pn_dc_scenario]

    st.info({
        "de":"KI generiert kompletten Entlassungsbrief: Diagnosen (ICD-10), Befunde, Inhalator-Reconciliation, O2-Rezept, schriftlicher Aktionsplan.",
        "en":"AI generates complete discharge letter: diagnoses (ICD-10), findings, inhaler reconciliation, O2 prescription, written action plan.",
        "ro":"AI genereaza scrisoare externare completa: diagnostice (ICD-10), rezultate, reconciliere inhalatoare, reteta O2, plan actiune.",
        "es":"La IA genera informe alta completo: diagnosticos (ICD-10), resultados, reconciliacion inhaladores, receta O2, plan de accion.",
        "ar":"     :  (ICD-10)      O2   ."
    }.get(lang,"AI generates complete discharge letter."))

    notes = _notes_and_audio(lang, "discharge", client, height=240,
        placeholder_extra=" (Diagnosen, Verlauf, Spirometrie, Inhalatoren, O2-Bedarf)")

    if st.button(tx("generate", lang), type="primary", use_container_width=True, key="pn_gen_dc"):
        if not notes.strip():
            st.warning(tx("fill_warn", lang))
        else:
            with st.spinner(tx("generating", lang)):
                try:
                    st.session_state.pn_res_discharge = _call_ai(
                        _p_discharge(notes, lang, country, style, scenario),
                        client, max_tokens=3000
                    )
                except Exception as e:
                    st.error(str(e))

    _output_block("pn_res_discharge", {
        "de":"Pneumologischer Entlassungsbrief","en":"Pulmonary Discharge Letter",
        "ro":"Scrisoare Externare Pneumologie","es":"Informe Alta Neumologia","ar":"  "
    }.get(lang,"Discharge Letter"), lang)


# ---------------------------------------------------------------------------
# MAIN RENDER
# ---------------------------------------------------------------------------
def render(lang="en", openai_client=None, supabase_client=None, user_code=""):
    """
    Main entry point.

    Standalone: streamlit run pneumology.py
    Module: import pneumology; pneumology.render(lang_code, client, supabase, user_code)
    """
    _init()

    # Load doctor style if supabase connected
    style_from_db = ""
    if supabase_client and user_code:
        try:
            resp = supabase_client.table("invite_codes").select("style_prompt").eq("code", user_code).execute()
            if resp.data:
                style_from_db = resp.data[0].get("style_prompt", "")
        except Exception:
            pass

    # -- STANDALONE MODE: full app with language + login ------------------
    standalone = openai_client is None

    if standalone:
        st.set_page_config(
            page_title="Lung Clinic OS",
            page_icon="?",
            layout="wide"
        )
        # Language selector
        lang_map = {"Deutsch":"de","English":"en","Romana":"ro","Espanol":"es","":"ar"}
        chosen = st.selectbox(
            "Language / Sprache / Limba / Idioma",
            list(lang_map.keys()),
            index=1
        )
        lang = lang_map[chosen]

        # API Key input for standalone mode
        api_key = st.text_input(
            "OpenAI API Key", type="password",
            help="Your key is not stored. Session only."
        )
        if api_key:
            openai_client = OpenAI(api_key=api_key)
        else:
            st.warning("Enter your OpenAI API key to generate reports.")

    # -- HEADER ----------------------------------------------------------
    st.markdown("## ? " + tx("app_title", lang) + " | " + tx("app_sub", lang))
    st.error(tx("disclaimer", lang))
    st.markdown("---")

    # -- COUNTRY + RESET --------------------------------------------------
    col_c, col_r = st.columns([5, 1])
    with col_c:
        c_opts = tx("country_opts", lang)
        st.session_state.pn_country = st.selectbox(
            tx("country", lang), range(len(c_opts)),
            format_func=lambda i: c_opts[i],
            index=st.session_state.pn_country, key="pn_sel_country"
        )
    with col_r:
        if st.button(tx("new", lang), key="pn_reset"):
            keys_to_clear = [k for k in st.session_state if k.startswith("pn_") or k.startswith("notes_pn") or k.startswith("audio_") or k.startswith("sb_")]
            for k in keys_to_clear:
                del st.session_state[k]
            st.rerun()

    country, base_style = _country_info(lang, st.session_state.pn_country)
    style = base_style + (" Doctor style: " + style_from_db if style_from_db else "")

    # -- 10 TABS ----------------------------------------------------------
    tabs = st.tabs([
        "? " + tx("t_intake", lang),
        "? " + tx("t_spiro", lang),
        "? " + tx("t_copd", lang),
        "? " + tx("t_imaging", lang),
        "? " + tx("t_nodule", lang),
        "? " + tx("t_infection", lang),
        "? " + tx("t_icu", lang),
        "? " + tx("t_sleep", lang),
        "? " + tx("t_followup", lang),
        "? " + tx("t_discharge", lang),
    ])

    with tabs[0]: _tab_intake(lang, openai_client, country, style)
    with tabs[1]: _tab_spirometry(lang, openai_client, country, style)
    with tabs[2]: _tab_copd(lang, openai_client, country, style)
    with tabs[3]: _tab_imaging(lang, openai_client, country, style)
    with tabs[4]: _tab_nodule(lang, openai_client, country, style)
    with tabs[5]: _tab_infection(lang, openai_client, country, style)
    with tabs[6]: _tab_icu(lang, openai_client, country, style)
    with tabs[7]: _tab_sleep(lang, openai_client, country, style)
    with tabs[8]: _tab_followup(lang, openai_client, country, style)
    with tabs[9]: _tab_discharge(lang, openai_client, country, style)

    # -- FOOTER ----------------------------------------------------------
    st.markdown("---")
    st.caption("? Lung Clinic OS | " + tx("disclaimer", lang))


# -- STANDALONE ENTRY POINT ----------------------------------------------
if __name__ == "__main__":
    render()
